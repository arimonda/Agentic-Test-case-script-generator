"""
Requirement Document Parser.

Extracts test intents and expected results from .docx and .xlsx files.
Uses AI to interpret unstructured business requirements into structured
test steps.
"""

import logging
from pathlib import Path
from typing import Optional

import openpyxl
from docx import Document as DocxDocument

from core.ai_engine import AIEngine
from models.schemas import (
    ActionType,
    ParsedRequirements,
    RequirementItem,
    TestCase,
    TestStepInput,
)

logger = logging.getLogger(__name__)

PARSE_PROMPT_TEMPLATE = """You are an expert QA analyst. Parse the following requirement text
into structured test steps.

## Requirement Text
{requirement_text}

## Instructions
Return ONLY valid JSON (no markdown fences) with this structure:

{{
    "test_case_name": "<descriptive name for this test case>",
    "description": "<what this test case validates>",
    "preconditions": "<any setup needed>",
    "steps": [
        {{
            "step_number": 1,
            "intent": "<what the user does, e.g. 'Click the Login button'>",
            "action_type": "<click|fill|select|check|uncheck|hover|navigate|wait|assert_visible|assert_text|assert_value|screenshot>",
            "input_data": "<data to enter, or null>",
            "expected_result": "<what should happen after this step>"
        }}
    ]
}}

Rules:
- Break complex requirements into atomic, sequential steps.
- Start with navigation if a URL is mentioned.
- Include assertion steps to verify expected outcomes.
- Use "fill" for text entry, "click" for button/link interaction.
- If a step involves entering data, put the value in "input_data".
"""


class RequirementParser:
    """Parses .docx and .xlsx requirement files into TestCase objects."""

    def __init__(self, ai_engine: AIEngine):
        self.ai = ai_engine

    def parse_file(self, file_path: str, target_url: str = "") -> list[TestCase]:
        """
        Parse a requirement file and return a list of TestCase objects.

        Supports .docx and .xlsx formats.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Requirement file not found: {file_path}")

        ext = path.suffix.lower()
        if ext == ".docx":
            raw_text = self._extract_docx(path)
        elif ext == ".xlsx":
            raw_text = self._extract_xlsx(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Use .docx or .xlsx")

        logger.info("Extracted %d chars from %s", len(raw_text), file_path)

        # Use AI to structure the requirements
        test_cases = self._ai_parse(raw_text, target_url, str(path))
        logger.info("Parsed %d test case(s) from %s", len(test_cases), file_path)
        return test_cases

    def parse_text(self, text: str, target_url: str = "") -> list[TestCase]:
        """Parse raw requirement text directly (for testing or piped input)."""
        return self._ai_parse(text, target_url, "direct_input")

    # --------------------------------------------------------------------- #
    # File Extractors
    # --------------------------------------------------------------------- #

    def _extract_docx(self, path: Path) -> str:
        """Extract all text content from a Word document."""
        doc = DocxDocument(str(path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)

    def _extract_xlsx(self, path: Path) -> str:
        """Extract content from all sheets of an Excel workbook."""
        wb = openpyxl.load_workbook(str(path), data_only=True)
        lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")

            headers = []
            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                cell_values = [str(c) if c is not None else "" for c in row]
                if row_idx == 1:
                    headers = cell_values
                    lines.append(" | ".join(headers))
                else:
                    row_data = []
                    for i, val in enumerate(cell_values):
                        header = headers[i] if i < len(headers) else f"Col{i}"
                        if val:
                            row_data.append(f"{header}: {val}")
                    if row_data:
                        lines.append("; ".join(row_data))

        return "\n".join(lines)

    # --------------------------------------------------------------------- #
    # AI-Powered Parsing
    # --------------------------------------------------------------------- #

    def _ai_parse(
        self, raw_text: str, target_url: str, source: str
    ) -> list[TestCase]:
        """Send raw text to AI for structured parsing."""
        # Split into chunks if text is very long
        chunks = self._chunk_text(raw_text, max_chars=8000)
        test_cases = []

        for i, chunk in enumerate(chunks):
            prompt = PARSE_PROMPT_TEMPLATE.format(requirement_text=chunk)

            try:
                data = self.ai.infer_json(prompt)
            except Exception as exc:
                logger.error("AI parsing failed for chunk %d: %s", i, exc)
                continue

            steps = []
            for step_data in data.get("steps", []):
                action_str = step_data.get("action_type", "click").lower()
                try:
                    action_type = ActionType(action_str)
                except ValueError:
                    action_type = ActionType.CLICK

                steps.append(TestStepInput(
                    step_number=step_data.get("step_number", len(steps) + 1),
                    intent=step_data["intent"],
                    action_type=action_type,
                    input_data=step_data.get("input_data"),
                    expected_result=step_data.get("expected_result", ""),
                ))

            tc = TestCase(
                name=data.get("test_case_name", f"Test Case {i + 1}"),
                description=data.get("description", ""),
                preconditions=data.get("preconditions", ""),
                target_url=target_url,
                steps=steps,
                source_file=source,
            )
            test_cases.append(tc)

        return test_cases

    def _chunk_text(self, text: str, max_chars: int = 8000) -> list[str]:
        """Split text into chunks that fit within AI context limits."""
        if len(text) <= max_chars:
            return [text]

        chunks = []
        lines = text.split("\n")
        current_chunk: list[str] = []
        current_len = 0

        for line in lines:
            if current_len + len(line) + 1 > max_chars and current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = []
                current_len = 0
            current_chunk.append(line)
            current_len += len(line) + 1

        if current_chunk:
            chunks.append("\n".join(current_chunk))

        return chunks
