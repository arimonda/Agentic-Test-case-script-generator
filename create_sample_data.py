"""
Sample Data & Documentation Generator.

Creates sample requirement documents (.docx, .xlsx), a User Guide,
and a User Manual for the Self-Healing Autonomous Test Agent.

Run once:  python create_sample_data.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side


def create_sample_docx_requirement():
    """Create a sample .docx requirement document targeting SauceDemo."""
    doc = Document()

    # Title
    title = doc.add_heading("Test Requirement: SauceDemo Login & Checkout Flow", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document specifies the functional test requirements for the "
        "SauceDemo e-commerce demo application (https://www.saucedemo.com). "
        "The agent should validate the login flow, product browsing, "
        "cart management, and checkout process."
    )

    # --- Test Scenario 1 ---
    doc.add_heading("Test Scenario 1: Valid User Login", level=1)
    doc.add_paragraph(
        "Verify that a user can successfully log in with valid credentials."
    )

    table = doc.add_table(rows=5, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Step", "Action", "Input Data", "Expected Result"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    steps = [
        ("1", "Navigate to the application URL", "https://www.saucedemo.com", "Login page is displayed with username and password fields"),
        ("2", "Enter the username", "standard_user", "Username field is populated"),
        ("3", "Enter the password", "secret_sauce", "Password field is populated (masked)"),
        ("4", "Click the Login button", "", "User is redirected to the Products/Inventory page"),
    ]
    for row_idx, (step, action, data, expected) in enumerate(steps, start=1):
        table.rows[row_idx].cells[0].text = step
        table.rows[row_idx].cells[1].text = action
        table.rows[row_idx].cells[2].text = data
        table.rows[row_idx].cells[3].text = expected

    doc.add_paragraph()

    # --- Test Scenario 2 ---
    doc.add_heading("Test Scenario 2: Add Product to Cart and Checkout", level=1)
    doc.add_paragraph(
        "After logging in, verify that the user can add a product to the cart "
        "and complete the checkout process."
    )

    table2 = doc.add_table(rows=8, cols=4)
    table2.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table2.rows[0].cells[i].text = h

    steps2 = [
        ("1", "Click on 'Sauce Labs Backpack' product title", "", "Product detail page is displayed"),
        ("2", "Click the 'Add to cart' button", "", "Button changes to 'Remove' and cart badge shows '1'"),
        ("3", "Click the shopping cart icon", "", "Cart page is displayed with the Backpack item"),
        ("4", "Click the 'Checkout' button", "", "Checkout information form is displayed"),
        ("5", "Enter first name", "John", "First name field is populated"),
        ("6", "Enter last name", "Doe", "Last name field is populated"),
        ("7", "Enter postal code", "12345", "Postal code field is populated"),
    ]
    for row_idx, (step, action, data, expected) in enumerate(steps2, start=1):
        table2.rows[row_idx].cells[0].text = step
        table2.rows[row_idx].cells[1].text = action
        table2.rows[row_idx].cells[2].text = data
        table2.rows[row_idx].cells[3].text = expected

    doc.add_paragraph()

    # --- Test Scenario 3 ---
    doc.add_heading("Test Scenario 3: Invalid Login Attempt", level=1)
    doc.add_paragraph(
        "Verify that the system displays an appropriate error message when "
        "invalid credentials are used."
    )

    table3 = doc.add_table(rows=5, cols=4)
    table3.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table3.rows[0].cells[i].text = h

    steps3 = [
        ("1", "Navigate to the application URL", "https://www.saucedemo.com", "Login page is displayed"),
        ("2", "Enter an invalid username", "invalid_user", "Username field is populated"),
        ("3", "Enter an invalid password", "wrong_pass", "Password field is populated"),
        ("4", "Click the Login button", "", "Error message is displayed: 'Username and password do not match'"),
    ]
    for row_idx, (step, action, data, expected) in enumerate(steps3, start=1):
        table3.rows[row_idx].cells[0].text = step
        table3.rows[row_idx].cells[1].text = action
        table3.rows[row_idx].cells[2].text = data
        table3.rows[row_idx].cells[3].text = expected

    path = Path("input/requirements/saucedemo_requirements.docx")
    doc.save(str(path))
    print(f"Created: {path}")


def create_sample_xlsx_requirement():
    """Create a sample .xlsx requirement document."""
    wb = openpyxl.Workbook()

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    # --- Sheet 1: Login Tests ---
    ws1 = wb.active
    ws1.title = "Login Tests"

    headers = ["Test ID", "Scenario", "Step", "Action", "Input Data", "Expected Result", "Priority"]
    for col, h in enumerate(headers, start=1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    login_data = [
        ["TC-001", "Valid Login", "1", "Open browser and navigate to https://www.saucedemo.com", "https://www.saucedemo.com", "Login page loads successfully", "High"],
        ["TC-001", "Valid Login", "2", "Enter username in the username field", "standard_user", "Username is entered", "High"],
        ["TC-001", "Valid Login", "3", "Enter password in the password field", "secret_sauce", "Password is entered (masked)", "High"],
        ["TC-001", "Valid Login", "4", "Click the Login button", "", "Redirected to inventory/products page", "High"],
        ["TC-002", "Locked Out User", "1", "Navigate to login page", "https://www.saucedemo.com", "Login page loads", "Medium"],
        ["TC-002", "Locked Out User", "2", "Enter locked out username", "locked_out_user", "Username is entered", "Medium"],
        ["TC-002", "Locked Out User", "3", "Enter password", "secret_sauce", "Password is entered", "Medium"],
        ["TC-002", "Locked Out User", "4", "Click Login", "", "Error: 'Sorry, this user has been locked out'", "Medium"],
    ]

    for row_idx, row_data in enumerate(login_data, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    for col in ws1.columns:
        ws1.column_dimensions[col[0].column_letter].width = 20

    # --- Sheet 2: Cart & Checkout Tests ---
    ws2 = wb.create_sheet("Cart & Checkout")

    for col, h in enumerate(headers, start=1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    cart_data = [
        ["TC-003", "Add to Cart", "1", "Login with standard_user / secret_sauce", "standard_user / secret_sauce", "Successfully logged in", "High"],
        ["TC-003", "Add to Cart", "2", "Click 'Add to cart' on Sauce Labs Backpack", "", "Cart badge shows 1", "High"],
        ["TC-003", "Add to Cart", "3", "Click 'Add to cart' on Sauce Labs Bike Light", "", "Cart badge shows 2", "High"],
        ["TC-003", "Add to Cart", "4", "Click the cart icon", "", "Cart page shows both items", "High"],
        ["TC-004", "Checkout Flow", "1", "From cart page, click Checkout", "", "Checkout info page appears", "High"],
        ["TC-004", "Checkout Flow", "2", "Enter first name", "Jane", "First name populated", "High"],
        ["TC-004", "Checkout Flow", "3", "Enter last name", "Smith", "Last name populated", "High"],
        ["TC-004", "Checkout Flow", "4", "Enter zip code", "90210", "Zip code populated", "High"],
        ["TC-004", "Checkout Flow", "5", "Click Continue", "", "Checkout overview page appears", "High"],
        ["TC-004", "Checkout Flow", "6", "Click Finish", "", "Order confirmation page displayed", "High"],
    ]

    for row_idx, row_data in enumerate(cart_data, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    for col in ws2.columns:
        ws2.column_dimensions[col[0].column_letter].width = 20

    path = Path("input/requirements/saucedemo_test_cases.xlsx")
    wb.save(str(path))
    print(f"Created: {path}")


def create_sample_plain_text_requirement():
    """Create a simple plain-text requirement for quick testing."""
    content = """Test Requirement: SauceDemo Quick Login Test

Application URL: https://www.saucedemo.com

Objective: Verify that a standard user can log in successfully.

Steps:
1. Open the browser and navigate to https://www.saucedemo.com
2. Type "standard_user" into the Username field
3. Type "secret_sauce" into the Password field
4. Click the "Login" button
5. Verify that the page title shows "Products" or the inventory page is displayed
6. Verify that at least one product item is visible on the page

Expected Outcome: User should see the product inventory page with items listed.

Test Data:
- Username: standard_user
- Password: secret_sauce
- Invalid Username: invalid_user
- Invalid Password: wrong_password
"""
    path = Path("input/requirements/quick_login_test.txt")
    path.write_text(content, encoding="utf-8")
    print(f"Created: {path}")


def create_saucedemo_manual():
    """Create a sample user manual for the SauceDemo app (used as RAG knowledge base)."""
    doc = Document()

    title = doc.add_heading("SauceDemo Application — User Manual", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This manual describes the SauceDemo e-commerce demo application "
        "available at https://www.saucedemo.com. It is used as a reference "
        "knowledge base for the AI Test Agent."
    )

    # --- Section: Login ---
    doc.add_heading("1. Login Page", level=1)
    doc.add_paragraph(
        "The login page is the entry point of the application. It contains:\n"
        "- A 'Username' text field (id: user-name)\n"
        "- A 'Password' text field (id: password)\n"
        "- A 'Login' button (id: login-button)\n\n"
        "Valid test users:\n"
        "- standard_user (normal access)\n"
        "- locked_out_user (will see an error)\n"
        "- problem_user (broken images)\n"
        "- performance_glitch_user (slow loading)\n"
        "- error_user (random errors)\n"
        "- visual_user (visual inconsistencies)\n\n"
        "All users share the password: secret_sauce\n\n"
        "If login fails, a red error banner appears below the password field "
        "with class 'error-message-container'."
    )

    # --- Section: Products/Inventory ---
    doc.add_heading("2. Products Page (Inventory)", level=1)
    doc.add_paragraph(
        "After successful login, users land on the Products page.\n\n"
        "Page structure:\n"
        "- Header bar with the app title 'Swag Labs' and a cart icon (class: shopping_cart_link)\n"
        "- A product sort dropdown (class: product_sort_container)\n"
        "- Product cards in a grid layout, each containing:\n"
        "  - Product image (class: inventory_item_img)\n"
        "  - Product name (class: inventory_item_name) — clickable for details\n"
        "  - Product description (class: inventory_item_desc)\n"
        "  - Product price (class: inventory_item_price)\n"
        "  - 'Add to cart' button (changes to 'Remove' after clicking)\n\n"
        "Available products:\n"
        "1. Sauce Labs Backpack — $29.99\n"
        "2. Sauce Labs Bike Light — $9.99\n"
        "3. Sauce Labs Bolt T-Shirt — $15.99\n"
        "4. Sauce Labs Fleece Jacket — $49.99\n"
        "5. Sauce Labs Onesie — $7.99\n"
        "6. Test.allTheThings() T-Shirt (Red) — $15.99\n\n"
        "The cart icon in the header shows a badge with the number of items in the cart."
    )

    # --- Section: Cart ---
    doc.add_heading("3. Shopping Cart", level=1)
    doc.add_paragraph(
        "Clicking the cart icon navigates to the cart page.\n\n"
        "Cart page elements:\n"
        "- Cart items list showing product name, description, price, and quantity\n"
        "- 'Remove' button next to each item\n"
        "- 'Continue Shopping' button — returns to the products page\n"
        "- 'Checkout' button — proceeds to checkout information\n\n"
        "The cart page URL is: https://www.saucedemo.com/cart.html"
    )

    # --- Section: Checkout ---
    doc.add_heading("4. Checkout Process", level=1)

    doc.add_heading("4.1 Checkout: Your Information", level=2)
    doc.add_paragraph(
        "This step collects shipping information:\n"
        "- First Name field (id: first-name)\n"
        "- Last Name field (id: last-name)\n"
        "- Zip/Postal Code field (id: postal-code)\n"
        "- 'Cancel' button — returns to cart\n"
        "- 'Continue' button — proceeds to checkout overview\n\n"
        "If any field is empty and Continue is clicked, an error message appears."
    )

    doc.add_heading("4.2 Checkout: Overview", level=2)
    doc.add_paragraph(
        "Shows a summary of the order:\n"
        "- List of items with prices\n"
        "- Payment Information section\n"
        "- Shipping Information section\n"
        "- Price Total (Item total + Tax)\n"
        "- 'Cancel' button — returns to products\n"
        "- 'Finish' button — completes the order"
    )

    doc.add_heading("4.3 Checkout: Complete", level=2)
    doc.add_paragraph(
        "After clicking Finish:\n"
        "- A 'Thank you for your order!' heading appears\n"
        "- A pony express image is displayed\n"
        "- 'Back Home' button returns to the products page\n\n"
        "URL: https://www.saucedemo.com/checkout-complete.html"
    )

    # --- Section: Sidebar Menu ---
    doc.add_heading("5. Sidebar Navigation Menu", level=1)
    doc.add_paragraph(
        "A hamburger menu (id: react-burger-menu-btn) is available on all pages after login.\n\n"
        "Menu items:\n"
        "- 'All Items' — navigates to the products page\n"
        "- 'About' — navigates to saucelabs.com\n"
        "- 'Logout' — logs out and returns to login page\n"
        "- 'Reset App State' — clears the cart and resets the app\n\n"
        "The close button for the menu has id: react-burger-cross-btn"
    )

    path = Path("input/manuals/saucedemo_user_manual.docx")
    doc.save(str(path))
    print(f"Created: {path}")


def create_user_guide():
    """Create the User Guide for the Self-Healing Autonomous Test Agent."""
    doc = Document()

    # --- Title Page ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading("Self-Healing Autonomous Test Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("USER GUIDE")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(47, 84, 150)
    run.bold = True

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 1.0 — February 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Getting Started",
        "   2.1 Prerequisites",
        "   2.2 Installation",
        "   2.3 Configuration",
        "3. Your First Test Run",
        "   3.1 Using a Requirement Document",
        "   3.2 Using Inline Text",
        "   3.3 Demo Mode",
        "4. Understanding the Output",
        "   4.1 Playwright Scripts",
        "   4.2 Word Test Cases",
        "   4.3 HTML Reports",
        "5. Working with the Knowledge Base (RAG)",
        "6. Switching AI Providers",
        "7. Self-Healing Explained",
        "8. Troubleshooting",
        "9. FAQ",
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # --- 1. Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The Self-Healing Autonomous Test Agent is an AI-powered automated testing "
        "ecosystem that reads your business requirements, opens a real browser, "
        "identifies UI elements using Computer Vision and DOM analysis, executes "
        "test steps, and automatically generates documentation and automation scripts."
    )
    doc.add_paragraph(
        "Key Capabilities:"
    )
    bullets = [
        "Reads .docx and .xlsx requirement documents",
        "Uses AI (Gemini or Claude) to identify UI elements from screenshots + HTML",
        "Self-healing locators — automatically retries with backup selectors if primary fails",
        "Generates Python Playwright test scripts (.py)",
        "Generates Word test case documents (.docx) with embedded screenshots",
        "Produces beautiful HTML execution reports with AI usage tracking",
        "RAG knowledge base for ingesting user manuals as domain context",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # --- 2. Getting Started ---
    doc.add_heading("2. Getting Started", level=1)

    doc.add_heading("2.1 Prerequisites", level=2)
    doc.add_paragraph("Before installing, ensure you have:")
    prereqs = [
        "Python 3.11 or higher installed",
        "An API key for Google Gemini and/or Anthropic Claude",
        "Internet connection (for AI API calls and browser testing)",
        "The target web application accessible from your machine",
    ]
    for p in prereqs:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("2.2 Installation", level=2)
    doc.add_paragraph("Follow these steps to install the agent:")
    install_steps = [
        "Open a terminal in the project directory",
        "Create a virtual environment: python -m venv .venv",
        "Activate it: .venv\\Scripts\\activate (Windows) or source .venv/bin/activate (Mac/Linux)",
        "Install dependencies: pip install -r requirements.txt",
        "Install Playwright browsers: playwright install chromium",
    ]
    for i, step in enumerate(install_steps, 1):
        doc.add_paragraph(f"Step {i}: {step}")

    doc.add_heading("2.3 Configuration", level=2)
    doc.add_paragraph(
        "Copy .env.example to .env and fill in your API keys:\n\n"
        "  GEMINI_API_KEY=your_key_here\n"
        "  CLAUDE_API_KEY=your_key_here\n"
        "  TARGET_URL=https://your-app.com\n\n"
        "The config/config.yaml file controls:\n"
        "- Which AI provider to use (GEMINI or CLAUDE)\n"
        "- Playwright browser settings (headless mode, viewport size, timeouts)\n"
        "- Self-healing retry count\n"
        "- Which artifacts to generate\n"
        "- RAG knowledge base settings"
    )

    # --- 3. Your First Test Run ---
    doc.add_heading("3. Your First Test Run", level=1)

    doc.add_heading("3.1 Using a Requirement Document", level=2)
    doc.add_paragraph(
        "Place your .docx or .xlsx requirement file in the input/requirements/ folder, "
        "then run:\n\n"
        "  python main.py --requirements input/requirements/saucedemo_requirements.docx "
        "--url https://www.saucedemo.com\n\n"
        "The agent will:\n"
        "1. Parse the requirement document using AI\n"
        "2. Open a Chromium browser\n"
        "3. Navigate to the target URL\n"
        "4. For each step: capture a screenshot + DOM, ask AI for locators, execute the action\n"
        "5. Generate a Playwright script, Word test case, and HTML report in output/"
    )

    doc.add_heading("3.2 Using Inline Text", level=2)
    doc.add_paragraph(
        "For quick tests, pass the requirement as text:\n\n"
        '  python main.py --text "Login with username standard_user and password '
        'secret_sauce, then verify the products page loads" --url https://www.saucedemo.com'
    )

    doc.add_heading("3.3 Demo Mode", level=2)
    doc.add_paragraph(
        "Demo mode executes a single action step — great for testing your setup:\n\n"
        '  python main.py --demo --url https://www.saucedemo.com --intent "Click the Login button"'
    )

    # --- 4. Understanding the Output ---
    doc.add_heading("4. Understanding the Output", level=1)
    doc.add_paragraph(
        "After execution, artifacts are saved in the output/ directory:"
    )

    doc.add_heading("4.1 Playwright Scripts (output/scripts/)", level=2)
    doc.add_paragraph(
        "Clean, runnable Python scripts with:\n"
        "- Browser setup and teardown\n"
        "- Each test step as a Playwright action\n"
        "- Primary locators as active code\n"
        "- Backup locators as comments for self-healing reference\n\n"
        "You can run the generated scripts directly:\n"
        "  python output/scripts/test_valid_login.py"
    )

    doc.add_heading("4.2 Word Test Cases (output/testcases/)", level=2)
    doc.add_paragraph(
        "Structured .docx documents containing:\n"
        "- Test case metadata (ID, description, preconditions, URL)\n"
        "- Step-by-step table with action, data, expected result, and status\n"
        "- Color-coded status (Green=Passed, Red=Failed, Yellow=Healed)\n"
        "- Embedded before/after screenshots for each step\n"
        "- Execution summary statistics"
    )

    doc.add_heading("4.3 HTML Reports (output/reports/)", level=2)
    doc.add_paragraph(
        "Modern dark-themed HTML dashboards with:\n"
        "- Pass/fail/heal statistics cards\n"
        "- Step-by-step cards with locator details\n"
        "- AI usage table (provider, tokens, latency, reasoning)\n"
        "- Also available as JSON for CI/CD integration"
    )

    # --- 5. RAG ---
    doc.add_heading("5. Working with the Knowledge Base (RAG)", level=1)
    doc.add_paragraph(
        "The RAG (Retrieval Augmented Generation) system lets you feed user manuals "
        "and technical documentation to the AI agent. This helps it understand:\n"
        "- Navigation paths not mentioned in requirements\n"
        "- Element IDs and CSS classes from technical docs\n"
        "- Business terminology and domain context\n\n"
        "To use:\n"
        "1. Place .docx, .txt, or .md files in input/manuals/\n"
        "2. Add --knowledge input/manuals/ to your command\n\n"
        "Example:\n"
        "  python main.py -r input/requirements/saucedemo_requirements.docx "
        "-u https://www.saucedemo.com -k input/manuals/"
    )

    # --- 6. Switching Providers ---
    doc.add_heading("6. Switching AI Providers", level=1)
    doc.add_paragraph(
        "You can switch between Gemini and Claude at any time:\n\n"
        "Via command line:\n"
        "  python main.py -r input/requirements/req.docx --provider GEMINI\n"
        "  python main.py -r input/requirements/req.docx --provider CLAUDE\n\n"
        "Via config.yaml:\n"
        '  ai_provider: "GEMINI"   # or "CLAUDE"\n\n'
        "Gemini is faster for real-time element identification.\n"
        "Claude is better for complex test planning and code generation."
    )

    # --- 7. Self-Healing ---
    doc.add_heading("7. Self-Healing Explained", level=1)
    doc.add_paragraph(
        "When the agent identifies a UI element, it creates three locators:\n\n"
        "1. PRIMARY — Most reliable (e.g., data-testid or id attribute)\n"
        "2. SECONDARY — Robust backup (e.g., ARIA label or CSS selector)\n"
        "3. TERTIARY — Last resort (e.g., XPath or visual text match)\n\n"
        "During execution:\n"
        "- The agent tries the PRIMARY locator first\n"
        "- If it fails (element not found, timeout), it retries with SECONDARY\n"
        "- If SECONDARY also fails, it tries TERTIARY\n"
        "- If all fail, the step is marked FAILED\n"
        "- If a backup succeeds, the step is marked HEALED\n\n"
        "This makes tests resilient to minor UI changes like renamed IDs or "
        "restructured HTML."
    )

    # --- 8. Troubleshooting ---
    doc.add_heading("8. Troubleshooting", level=1)
    problems = [
        ("'No API keys configured' error",
         "Ensure your .env file has GEMINI_API_KEY or CLAUDE_API_KEY set correctly."),
        ("Browser doesn't open",
         "Run 'playwright install chromium' to install browser binaries."),
        ("AI returns garbled locators",
         "Try switching to a different provider or increasing max_tokens in config.yaml."),
        ("Steps keep failing",
         "Check if the target URL is correct and accessible. Look at the screenshots in output/screenshots/ to see what the browser saw."),
        ("'Module not found' errors",
         "Ensure your virtual environment is activated and all deps are installed: pip install -r requirements.txt"),
    ]
    for problem, solution in problems:
        doc.add_paragraph(f"Problem: {problem}", style="List Bullet")
        doc.add_paragraph(f"Solution: {solution}")
        doc.add_paragraph()

    # --- 9. FAQ ---
    doc.add_heading("9. FAQ", level=1)
    faqs = [
        ("Can I use both Gemini and Claude in the same run?",
         "Currently, each run uses one provider. You can run the same requirements twice with different providers for comparison."),
        ("Does the agent modify my application?",
         "No. The agent only reads the DOM and takes screenshots. It performs UI actions (click, type) but does not alter your backend or database."),
        ("Can I edit the generated scripts?",
         "Absolutely. The generated .py files are designed to be clean and editable. Use them as a starting point and customize as needed."),
        ("How much does it cost per run?",
         "Cost depends on the AI provider. Each step sends a screenshot + HTML (~1000-3000 tokens). A 10-step test typically costs a few cents with Gemini."),
    ]
    for q, a in faqs:
        doc.add_heading(q, level=3)
        doc.add_paragraph(a)

    path = Path("docs/User_Guide.docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"Created: {path}")


def create_user_manual():
    """Create the technical User Manual / Reference Manual."""
    doc = Document()

    # --- Title ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_heading("Self-Healing Autonomous Test Agent", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("TECHNICAL REFERENCE MANUAL")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(47, 84, 150)
    run.bold = True

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("Version 1.0 — February 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # --- TOC ---
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. System Architecture",
        "2. Module Reference",
        "   2.1 core/agent.py — AutonomousTestAgent",
        "   2.2 core/ai_engine.py — AIEngine",
        "   2.3 core/locator_engine.py — LocatorEngine",
        "   2.4 core/action_executor.py — ActionExecutor",
        "   2.5 core/state_capture.py — StateCaptureEngine",
        "   2.6 core/requirement_parser.py — RequirementParser",
        "   2.7 core/config_loader.py — load_config()",
        "   2.8 generators/ — Artifact Generators",
        "   2.9 knowledge/rag_engine.py — RAGEngine",
        "3. Data Models Reference",
        "4. Configuration Reference",
        "5. CLI Reference",
        "6. API Integration Details",
        "7. Extending the Agent",
        "8. Security Considerations",
    ]
    for item in toc:
        doc.add_paragraph(item)

    doc.add_page_break()

    # --- 1. Architecture ---
    doc.add_heading("1. System Architecture", level=1)
    doc.add_paragraph(
        "The system follows a pipeline architecture with these stages:\n\n"
        "INGESTION → STATE CAPTURE → AI INFERENCE → ACTION EXECUTION → ARTIFACT GENERATION\n\n"
        "Data Flow:\n"
        "1. RequirementParser reads .docx/.xlsx and uses AI to extract structured TestCase objects\n"
        "2. RAGEngine (optional) ingests user manuals into a ChromaDB vector store\n"
        "3. For each TestStep:\n"
        "   a. StateCaptureEngine captures screenshot + DOM via Playwright\n"
        "   b. LocatorEngine sends screenshot + minified HTML to AI for element identification\n"
        "   c. ActionExecutor executes the Playwright action with self-healing retry\n"
        "4. ScriptGenerator, DocxGenerator, ReportGenerator run in parallel threads\n\n"
        "Thread Model:\n"
        "- Main thread: async Playwright event loop (asyncio)\n"
        "- Background ThreadPoolExecutor (3 workers): artifact generation"
    )

    # --- 2. Module Reference ---
    doc.add_heading("2. Module Reference", level=1)

    doc.add_heading("2.1 core/agent.py — AutonomousTestAgent", level=2)
    doc.add_paragraph(
        "The central orchestrator class.\n\n"
        "Constructor: AutonomousTestAgent(config: AppConfig)\n"
        "  - Initializes all sub-engines based on configuration\n\n"
        "Methods:\n"
        "  async run(requirement_file, requirement_text, target_url, knowledge_dir) → list[TestCaseResult]\n"
        "    Full pipeline execution. Parses requirements, opens browser, executes all steps, generates artifacts.\n\n"
        "  async run_single_step(target_url, intent, action_type, input_data) → TestStepResult\n"
        "    Execute a single step for demo/testing purposes."
    )

    doc.add_heading("2.2 core/ai_engine.py — AIEngine", level=2)
    doc.add_paragraph(
        "Unified multi-model inference engine.\n\n"
        "Constructor: AIEngine(config: AppConfig)\n\n"
        "Methods:\n"
        "  infer(prompt, images, provider, temperature, max_tokens) → AIResponse\n"
        "    Send a text+image prompt to the selected AI provider.\n\n"
        "  infer_json(prompt, images, provider) → dict\n"
        "    Convenience method that parses the AI response as JSON.\n\n"
        "Supported Providers:\n"
        "  - GEMINI: Uses google.generativeai SDK with native multimodal input\n"
        "  - CLAUDE: Uses anthropic SDK with base64 image messages"
    )

    doc.add_heading("2.3 core/locator_engine.py — LocatorEngine", level=2)
    doc.add_paragraph(
        "AI-driven element identification.\n\n"
        "Methods:\n"
        "  identify_locators(page_state, step, provider) → LocatorSet\n"
        "    Sends screenshot + minified HTML to AI with a structured prompt.\n"
        "    Returns a LocatorSet with primary/secondary/tertiary locators.\n\n"
        "  identify_multiple_locators(page_state, steps) → list[LocatorSet]\n"
        "    Batch identification for multiple steps on the same page.\n\n"
        "Prompt Engineering:\n"
        "  The locator prompt instructs the AI to:\n"
        "  1. Study the screenshot visually\n"
        "  2. Cross-reference with HTML\n"
        "  3. Return JSON with three ranked locators\n"
        "  4. Prioritize: data-testid > id > aria-label > CSS > XPath > visual text"
    )

    doc.add_heading("2.4 core/action_executor.py — ActionExecutor", level=2)
    doc.add_paragraph(
        "Self-healing action execution.\n\n"
        "Methods:\n"
        "  async execute_step(page, step, locators) → TestStepResult\n\n"
        "Self-Healing Logic:\n"
        "  1. Try PRIMARY locator → if success, status = PASSED\n"
        "  2. Try SECONDARY locator → if success, status = HEALED\n"
        "  3. Try TERTIARY locator → if success, status = HEALED\n"
        "  4. All failed → status = FAILED\n\n"
        "Supported Actions:\n"
        "  click, fill, select, check, uncheck, hover, navigate, wait,\n"
        "  assert_visible, assert_text, assert_value, screenshot"
    )

    doc.add_heading("2.5 core/state_capture.py — StateCaptureEngine", level=2)
    doc.add_paragraph(
        "Captures page state for AI analysis.\n\n"
        "Methods:\n"
        "  async capture(page, step_label, full_page) → PageState\n"
        "    Returns: URL, title, screenshot (file + base64), raw DOM, minified HTML, visible text\n\n"
        "  async capture_element(page, selector, label) → bytes | None\n"
        "    Capture a screenshot of a single element."
    )

    doc.add_heading("2.6 core/requirement_parser.py — RequirementParser", level=2)
    doc.add_paragraph(
        "Document parsing engine.\n\n"
        "Methods:\n"
        "  parse_file(file_path, target_url) → list[TestCase]\n"
        "    Reads .docx or .xlsx, extracts text, sends to AI for structured parsing.\n\n"
        "  parse_text(text, target_url) → list[TestCase]\n"
        "    Parse raw text directly.\n\n"
        "Document Extraction:\n"
        "  - .docx: Reads paragraphs and table cells\n"
        "  - .xlsx: Reads all sheets, preserving headers and row structure\n\n"
        "AI Parsing:\n"
        "  The extracted text is sent to the AI with a prompt that asks for:\n"
        "  - test_case_name, description, preconditions\n"
        "  - Steps with: step_number, intent, action_type, input_data, expected_result"
    )

    doc.add_heading("2.7 core/config_loader.py — load_config()", level=2)
    doc.add_paragraph(
        "Configuration loading function.\n\n"
        "  load_config(config_path, env_path) → AppConfig\n\n"
        "Precedence: Environment variables > YAML values > Defaults\n\n"
        "Reads from:\n"
        "  1. config/config.yaml — YAML configuration\n"
        "  2. .env — Environment variables (via python-dotenv)"
    )

    doc.add_heading("2.8 generators/ — Artifact Generators", level=2)
    doc.add_paragraph(
        "Three parallel generators:\n\n"
        "ScriptGenerator (script_generator.py)\n"
        "  generate(test_result) → str (file path)\n"
        "  Produces PEP8-compliant .py files with setup, teardown, and locator comments.\n\n"
        "DocxGenerator (docx_generator.py)\n"
        "  generate(test_result) → str (file path)\n"
        "  Produces .docx files with metadata table, step table, screenshots, summary.\n\n"
        "ReportGenerator (report_generator.py)\n"
        "  log_ai_usage(step_number, provider, tokens, latency_ms, reasoning)\n"
        "  generate(test_result) → str (file path)\n"
        "  Produces dark-themed HTML reports + JSON reports."
    )

    doc.add_heading("2.9 knowledge/rag_engine.py — RAGEngine", level=2)
    doc.add_paragraph(
        "RAG knowledge base engine.\n\n"
        "Methods:\n"
        "  ingest_document(file_path) → int (chunks created)\n"
        "  ingest_directory(dir_path) → int (total chunks)\n"
        "  query(question, n_results) → list[dict]\n"
        "  get_context_for_step(intent, max_chars) → str\n"
        "  clear() → None\n\n"
        "Storage: ChromaDB with cosine similarity\n"
        "Chunking: RecursiveCharacterTextSplitter (1000 chars, 200 overlap)\n"
        "Supported formats: .docx, .txt, .md"
    )

    # --- 3. Data Models ---
    doc.add_heading("3. Data Models Reference", level=1)
    doc.add_paragraph(
        "All models are defined in models/schemas.py using Pydantic.\n\n"
        "Core Models:\n"
        "  - Locator: strategy, value, confidence, description\n"
        "  - LocatorSet: element_name, primary, secondary, tertiary\n"
        "  - TestStepInput: step_number, intent, action_type, input_data, expected_result\n"
        "  - TestStepResult: step_input, status, locators_used, screenshots, timing\n"
        "  - TestCase: name, description, steps, target_url, preconditions\n"
        "  - TestCaseResult: test_case, step_results, overall_status, timing\n"
        "  - PageState: url, title, screenshot, dom_html, minified_html\n"
        "  - AIRequest / AIResponse: provider, model, content, usage\n"
        "  - AppConfig: all runtime configuration\n\n"
        "Enums:\n"
        "  - LocatorStrategy: test_id, id, aria, css, xpath, visual\n"
        "  - ActionType: click, fill, select, check, uncheck, hover, navigate, wait, assert_*\n"
        "  - StepStatus: pending, running, passed, failed, healed, skipped\n"
        "  - AIProvider: GEMINI, CLAUDE"
    )

    # --- 4. Config Reference ---
    doc.add_heading("4. Configuration Reference", level=1)

    config_table = doc.add_table(rows=1, cols=4)
    config_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Setting", "YAML Path", "Env Variable", "Default"]):
        config_table.rows[0].cells[i].text = h

    configs = [
        ("AI Provider", "ai_provider", "AI_PROVIDER", "GEMINI"),
        ("Gemini Key", "api_keys.gemini", "GEMINI_API_KEY", "—"),
        ("Claude Key", "api_keys.claude", "CLAUDE_API_KEY", "—"),
        ("Target URL", "—", "TARGET_URL", "—"),
        ("Headless", "playwright.headless", "—", "false"),
        ("Browser", "playwright.browser", "—", "chromium"),
        ("Viewport Width", "playwright.viewport.width", "—", "1920"),
        ("Viewport Height", "playwright.viewport.height", "—", "1080"),
        ("Timeout (ms)", "playwright.timeout", "—", "30000"),
        ("Max Retries", "self_healing.max_retries", "—", "3"),
        ("Output Dir", "artifacts.output_dir", "—", "output"),
        ("Generate Script", "artifacts.generate_script", "—", "true"),
        ("Generate DOCX", "artifacts.generate_docx", "—", "true"),
        ("Generate Report", "artifacts.generate_report", "—", "true"),
        ("RAG Enabled", "rag.enabled", "—", "true"),
        ("Chunk Size", "rag.chunk_size", "—", "1000"),
    ]
    for row_data in configs:
        row = config_table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    # --- 5. CLI Reference ---
    doc.add_heading("5. CLI Reference", level=1)
    doc.add_paragraph(
        "Usage: python main.py [OPTIONS]\n\n"
        "Options:\n"
        "  --requirements, -r  Path to .docx/.xlsx requirement file\n"
        "  --text, -t          Inline requirement text\n"
        "  --url, -u           Target application URL\n"
        "  --knowledge, -k     Directory with user manuals for RAG\n"
        "  --config, -c        Path to config YAML (default: config/config.yaml)\n"
        "  --demo              Run in demo mode (single step)\n"
        "  --intent            Intent for demo mode\n"
        "  --provider          Override AI provider (GEMINI or CLAUDE)\n"
        "  --headless          Run browser in headless mode\n"
        "  --log-level         Logging level (DEBUG, INFO, WARNING, ERROR)\n\n"
        "Examples:\n"
        "  python main.py -r input/requirements/req.docx -u https://app.com\n"
        "  python main.py -t \"Login and verify dashboard\" -u https://app.com\n"
        "  python main.py --demo -u https://app.com --intent \"Click Login\"\n"
        "  python main.py -r req.docx -u https://app.com -k input/manuals/ --provider GEMINI"
    )

    # --- 6. API Integration ---
    doc.add_heading("6. API Integration Details", level=1)
    doc.add_paragraph(
        "Gemini Integration:\n"
        "  SDK: google-generativeai\n"
        "  Model: gemini-2.5-flash (configurable)\n"
        "  Input: Multi-part content (images as bytes + text prompt)\n"
        "  Authentication: API key via GEMINI_API_KEY environment variable\n\n"
        "Claude Integration:\n"
        "  SDK: anthropic\n"
        "  Model: claude-sonnet-4-20250514 (configurable)\n"
        "  Input: Messages API with base64 image content blocks\n"
        "  Authentication: API key via CLAUDE_API_KEY environment variable\n\n"
        "Token Usage:\n"
        "  Each step sends: ~500-2000 tokens for minified HTML + ~1000 tokens for the prompt\n"
        "  Images: Counted separately by each provider\n"
        "  Total per step: approximately 2000-4000 tokens"
    )

    # --- 7. Extending ---
    doc.add_heading("7. Extending the Agent", level=1)
    doc.add_paragraph(
        "Adding a New AI Provider:\n"
        "  1. Add the provider to AIProvider enum in models/schemas.py\n"
        "  2. Add a _infer_<provider>() method in core/ai_engine.py\n"
        "  3. Add model configuration in config.yaml\n"
        "  4. Add API key support in .env and config_loader.py\n\n"
        "Adding a New Action Type:\n"
        "  1. Add the action to ActionType enum in models/schemas.py\n"
        "  2. Add the case handler in ActionExecutor._perform_action()\n"
        "  3. Add the code generation in ScriptGenerator._build_action_code()\n\n"
        "Adding a New Locator Strategy:\n"
        "  1. Add the strategy to LocatorStrategy enum in models/schemas.py\n"
        "  2. Add the Playwright resolution in ActionExecutor._resolve_locator()\n"
        "  3. Add the code generation in Locator.to_playwright()\n\n"
        "Custom Report Templates:\n"
        "  Edit the HTML_TEMPLATE constant in generators/report_generator.py"
    )

    # --- 8. Security ---
    doc.add_heading("8. Security Considerations", level=1)
    doc.add_paragraph(
        "API Key Management:\n"
        "  - NEVER commit .env files to version control\n"
        "  - The .gitignore is pre-configured to exclude .env\n"
        "  - Use environment variables in CI/CD pipelines\n"
        "  - Rotate keys periodically\n\n"
        "Data Privacy:\n"
        "  - Screenshots and DOM content are sent to external AI APIs\n"
        "  - Do not test applications with sensitive/PII data without approval\n"
        "  - Generated artifacts may contain screenshots of the application\n"
        "  - RAG vector store is local (ChromaDB on disk) — no external storage\n\n"
        "Network:\n"
        "  - The agent makes outbound HTTPS calls to AI APIs\n"
        "  - Playwright connects to the target URL over the network\n"
        "  - No inbound connections are opened"
    )

    path = Path("docs/Technical_Reference_Manual.docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"Created: {path}")


if __name__ == "__main__":
    print("Creating sample data and documentation...\n")

    print("--- Sample Requirements ---")
    create_sample_docx_requirement()
    create_sample_xlsx_requirement()
    create_sample_plain_text_requirement()

    print("\n--- Knowledge Base ---")
    create_saucedemo_manual()

    print("\n--- Documentation ---")
    create_user_guide()
    create_user_manual()

    print("\nAll files created successfully!")
